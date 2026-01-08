"""
Document parser for extracting text from PDF, DOCX, and XLSX files.
"""
import os
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import load_workbook


def allowed_file(filename: str, allowed_extensions: set) -> bool:
    """Check if file has an allowed extension."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed_extensions


def get_file_extension(filename: str) -> str:
    """Get the file extension in lowercase."""
    return filename.rsplit(".", 1)[1].lower() if "." in filename else ""


def parse_pdf(file_path: str) -> dict:
    """
    Extract text from a PDF file.

    Returns:
        dict with 'text' (full text), 'pages' (list of page texts), 'metadata'
    """
    reader = PdfReader(file_path)
    pages = []
    full_text = []

    for page_num, page in enumerate(reader.pages, 1):
        page_text = page.extract_text() or ""
        pages.append({
            "page_number": page_num,
            "text": page_text
        })
        full_text.append(page_text)

    metadata = {}
    if reader.metadata:
        metadata = {
            "title": reader.metadata.get("/Title", ""),
            "author": reader.metadata.get("/Author", ""),
            "subject": reader.metadata.get("/Subject", ""),
            "creator": reader.metadata.get("/Creator", ""),
        }

    return {
        "text": "\n\n".join(full_text),
        "pages": pages,
        "page_count": len(pages),
        "metadata": metadata,
        "format": "pdf"
    }


def parse_docx(file_path: str) -> dict:
    """
    Extract text from a Word document.

    Returns:
        dict with 'text' (full text), 'paragraphs' (list), 'tables' (list), 'metadata'
    """
    doc = Document(file_path)

    paragraphs = []
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else "Normal"
            })
            full_text.append(para.text)

    # Extract tables
    tables = []
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append({
            "table_number": table_idx + 1,
            "data": table_data
        })

    # Get metadata from core properties
    metadata = {}
    try:
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
        }
    except Exception:
        pass

    return {
        "text": "\n\n".join(full_text),
        "paragraphs": paragraphs,
        "tables": tables,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "metadata": metadata,
        "format": "docx"
    }


def parse_xlsx(file_path: str) -> dict:
    """
    Extract text from an Excel file.

    Returns:
        dict with 'text' (full text), 'sheets' (list of sheet data), 'metadata'
    """
    workbook = load_workbook(file_path, data_only=True)

    sheets = []
    full_text = []

    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        sheet_data = []
        sheet_text = []

        for row in sheet.iter_rows():
            row_data = []
            row_text = []
            for cell in row:
                value = str(cell.value) if cell.value is not None else ""
                row_data.append(value)
                if value:
                    row_text.append(value)
            sheet_data.append(row_data)
            if row_text:
                sheet_text.append(" | ".join(row_text))

        sheets.append({
            "name": sheet_name,
            "data": sheet_data,
            "row_count": len(sheet_data),
            "col_count": sheet.max_column
        })

        if sheet_text:
            full_text.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(sheet_text))

    return {
        "text": "\n\n".join(full_text),
        "sheets": sheets,
        "sheet_count": len(sheets),
        "metadata": {},
        "format": "xlsx"
    }


def parse_document(file_path: str) -> dict:
    """
    Parse a document and extract its text content.

    Args:
        file_path: Path to the document file

    Returns:
        dict containing extracted text and metadata

    Raises:
        ValueError: If file format is not supported
        FileNotFoundError: If file does not exist
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    extension = get_file_extension(file_path)

    if extension == "pdf":
        return parse_pdf(file_path)
    elif extension == "docx":
        return parse_docx(file_path)
    elif extension == "xlsx":
        return parse_xlsx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {extension}")


def extract_sections(text: str) -> list:
    """
    Attempt to identify section headers and their content from text.
    This is a heuristic approach that looks for common patterns.

    Returns:
        List of dicts with 'header' and 'content' keys
    """
    import re

    sections = []

    # Common section patterns
    patterns = [
        r'^(\d+\.)\s+(.+?)(?=\n\d+\.|$)',  # 1. Section Name
        r'^(\d+\.\d+)\s+(.+?)(?=\n\d+\.\d+|$)',  # 1.1 Subsection
        r'^(Section\s+\d+[.:])(.+?)(?=Section\s+\d+|$)',  # Section 1:
        r'^(ARTICLE\s+[IVX\d]+[.:])(.+?)(?=ARTICLE\s+|$)',  # ARTICLE I:
        r'^([A-Z][A-Z\s]{2,}:)(.+?)(?=[A-Z][A-Z\s]{2,}:|$)',  # HEADER:
    ]

    for pattern in patterns:
        matches = re.findall(pattern, text, re.MULTILINE | re.DOTALL | re.IGNORECASE)
        if matches:
            for match in matches:
                sections.append({
                    "header": match[0].strip(),
                    "content": match[1].strip() if len(match) > 1 else ""
                })
            break  # Use first pattern that matches

    # If no patterns matched, try to split by common separators
    if not sections:
        lines = text.split("\n")
        current_section = {"header": "Document Content", "content": ""}

        for line in lines:
            line = line.strip()
            # Check if line looks like a header (short, possibly numbered, possibly all caps)
            if line and len(line) < 100 and (
                re.match(r'^\d+\.', line) or
                line.isupper() or
                re.match(r'^[A-Z][a-z]+ [A-Z][a-z]+', line)
            ):
                if current_section["content"]:
                    sections.append(current_section)
                current_section = {"header": line, "content": ""}
            else:
                current_section["content"] += line + "\n"

        if current_section["content"]:
            sections.append(current_section)

    return sections
